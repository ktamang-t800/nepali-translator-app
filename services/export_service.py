import re
from datetime import datetime
from pathlib import Path
from typing import Dict, List

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt

from app.config import OUTPUTS_DIR


def sanitize_output_name(file_name: str) -> str:
    """
    Make a safe file name for the exported Word document.
    """
    clean_name = file_name.strip().replace(" ", "_")
    clean_name = re.sub(r"[^A-Za-z0-9._-]", "", clean_name)

    if not clean_name:
        clean_name = "translated_document"

    return clean_name


def combine_translated_chunks(translated_chunk_items: List[Dict]) -> str:
    """
    Combine translated chunks in order into one text body.
    """
    ordered_chunks = sorted(
        translated_chunk_items,
        key=lambda item: int(item.get("chunk_number", 0))
    )

    combined_parts: List[str] = []

    for chunk in ordered_chunks:
        translated_text = (chunk.get("translated_text", "") or "").strip()
        if translated_text:
            combined_parts.append(translated_text)

    return "\n\n".join(combined_parts).strip()


def build_output_docx_path(source_file_name: str) -> Path:
    """
    Build the timestamped output .docx file path.
    """
    safe_name = sanitize_output_name(source_file_name)
    source_stem = Path(safe_name).stem
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_name = "{0}_{1}_translated.docx".format(source_stem, timestamp)
    return OUTPUTS_DIR / output_name


def looks_like_page_label(text: str) -> bool:
    """
    Detect lines like [Page 1].
    """
    return bool(re.match(r"^\[Page\s+\d+\]$", text.strip(), flags=re.IGNORECASE))


def looks_like_heading(text: str) -> bool:
    """
    Simple heading detection for cleaner Word formatting.
    """
    clean_text = text.strip()

    if not clean_text:
        return False

    if len(clean_text) > 100:
        return False

    if clean_text.endswith("."):
        return False

    if clean_text.startswith("- ") or clean_text.startswith("* "):
        return False

    if re.match(r"^\d+\.\s+", clean_text):
        return False

    uppercase_ratio = sum(1 for ch in clean_text if ch.isupper()) / max(
        1, sum(1 for ch in clean_text if ch.isalpha())
    )

    if uppercase_ratio > 0.6 and len(clean_text.split()) <= 12:
        return True

    title_like = clean_text.istitle() and len(clean_text.split()) <= 12
    return title_like


def looks_like_bullet(text: str) -> bool:
    """
    Detect simple bullet-style lines.
    """
    clean_text = text.strip()

    if clean_text.startswith("- "):
        return True

    if clean_text.startswith("* "):
        return True

    if clean_text.startswith("• "):
        return True

    return False


def looks_like_numbered_item(text: str) -> bool:
    """
    Detect simple numbered-list lines.
    """
    clean_text = text.strip()
    return bool(re.match(r"^\d+\.\s+", clean_text))


def add_page_label(document: Document, text: str) -> None:
    """
    Add page label with stronger formatting and spacing.
    """
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text.strip())
    run.bold = True
    run.font.size = Pt(11)

    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(4)


def add_heading_paragraph(document: Document, text: str) -> None:
    """
    Add a heading-like paragraph.
    """
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text.strip())
    run.bold = True
    run.font.size = Pt(12)

    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(4)


def add_bullet_paragraph(document: Document, text: str) -> None:
    """
    Add a bullet-like paragraph.
    """
    clean_text = text.strip()

    if clean_text.startswith("- ") or clean_text.startswith("* ") or clean_text.startswith("• "):
        clean_text = clean_text[2:].strip()

    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.add_run(clean_text)
    paragraph.paragraph_format.space_after = Pt(2)


def add_numbered_paragraph(document: Document, text: str) -> None:
    """
    Add a numbered-like paragraph.
    """
    paragraph = document.add_paragraph(style="List Number")
    paragraph.add_run(text.strip())
    paragraph.paragraph_format.space_after = Pt(2)


def add_normal_paragraph(document: Document, text: str) -> None:
    """
    Add a standard body paragraph.
    """
    paragraph = document.add_paragraph(text.strip())
    paragraph.paragraph_format.space_after = Pt(6)


def add_formatted_content(document: Document, combined_text: str) -> None:
    """
    Add translated content with practical formatting improvements.
    """
    if not combined_text.strip():
        add_normal_paragraph(document, "[No translated content available]")
        return

    lines = combined_text.split("\n")

    for line in lines:
        clean_line = line.strip()

        if not clean_line:
            document.add_paragraph("")
            continue

        if looks_like_page_label(clean_line):
            add_page_label(document, clean_line)
        elif looks_like_bullet(clean_line):
            add_bullet_paragraph(document, clean_line)
        elif looks_like_numbered_item(clean_line):
            add_numbered_paragraph(document, clean_line)
        elif looks_like_heading(clean_line):
            add_heading_paragraph(document, clean_line)
        else:
            add_normal_paragraph(document, clean_line)


def export_translated_docx(
    source_file_name: str,
    translated_chunk_items: List[Dict],
    start_page: int,
    end_page: int
) -> Path:
    """
    Export translated chunks into a better formatted Word document.
    """
    document = Document()

    title = document.add_heading("Nepali to English Translation", level=0)
    title.alignment = 0

    meta_1 = document.add_paragraph()
    meta_1.add_run("Source file: ").bold = True
    meta_1.add_run("{0}".format(source_file_name))

    meta_2 = document.add_paragraph()
    meta_2.add_run("Translated page range: ").bold = True
    meta_2.add_run("Page {0} to Page {1}".format(start_page, end_page))

    meta_3 = document.add_paragraph()
    meta_3.add_run("Generated: ").bold = True
    meta_3.add_run(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))

    separator = document.add_paragraph()
    separator.add_run("")

    combined_text = combine_translated_chunks(translated_chunk_items)
    add_formatted_content(document, combined_text)

    output_path = build_output_docx_path(source_file_name)
    document.save(str(output_path))

    return output_path