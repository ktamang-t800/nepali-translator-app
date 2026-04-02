import json
from pathlib import Path
from typing import List

from docx import Document as DocxDocument

from app.config import EXTRACTED_DIR
from models.schemas import ExtractedDocument, ExtractedPage


def extract_docx_pages(docx_path: str, blocks_per_virtual_page: int = 12) -> ExtractedDocument:
    """
    Extract DOCX content and group it into virtual pages.

    DOCX files do not have reliable real page objects like PDFs,
    so we create virtual pages from paragraph/table-like blocks.
    """
    source_path = Path(docx_path)
    document = DocxDocument(docx_path)

    blocks: List[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)

    for table in document.tables:
        for row in table.rows:
            cell_texts: List[str] = []
            for cell in row.cells:
                clean_cell = cell.text.strip()
                if clean_cell:
                    cell_texts.append(clean_cell)

            if cell_texts:
                blocks.append(" | ".join(cell_texts))

    if not blocks:
        blocks = ["[No extractable text found in this DOCX document]"]

    pages: List[ExtractedPage] = []
    current_page_number = 1

    for index in range(0, len(blocks), blocks_per_virtual_page):
        page_blocks = blocks[index:index + blocks_per_virtual_page]
        page_text = "\n\n".join(page_blocks).strip()

        pages.append(
            ExtractedPage(
                page_number=current_page_number,
                text=page_text,
            )
        )
        current_page_number += 1

    return ExtractedDocument(
        source_file_name=source_path.name,
        source_file_path=str(source_path),
        file_type=".docx",
        total_pages=len(pages),
        pages=pages,
    )


def save_extracted_docx_result(document: ExtractedDocument) -> Path:
    """
    Save extracted DOCX text into a JSON file for later processing.
    """
    safe_stem = Path(document.source_file_name).stem
    output_name = "{0}_extracted.json".format(safe_stem)
    output_path = EXTRACTED_DIR / output_name

    with open(output_path, "w", encoding="utf-8") as json_file:
        json.dump(document.to_dict(), json_file, ensure_ascii=False, indent=2)

    return output_path


def build_docx_preview_lines(
    document: ExtractedDocument, max_pages: int = 2, max_chars: int = 500
) -> List[str]:
    """
    Build a small preview of extracted DOCX virtual pages for the UI.
    """
    preview_lines: List[str] = []

    for page in document.pages[:max_pages]:
        text = page.text.strip()

        if not text:
            text = "[No extractable text found in this virtual page]"

        if len(text) > max_chars:
            text = text[:max_chars] + "..."

        preview_lines.append("Virtual page {0}: {1}".format(page.page_number, text))

    return preview_lines